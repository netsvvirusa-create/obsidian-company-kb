"""Генерация .docx спецификации к договору из шаблона и данных Obsidian vault.

Расширяет функционал generate_contract.py: поддерживает вложенные переменные,
динамические таблицы услуг, расчёт итогов и НДС.
"""

import argparse
import copy
import json
import logging
import re
import sys
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.table import Table, _Row

from scripts.generate_contract import (
    VAR_PATTERN,
    build_contract_variables,
    check_remaining_vars,
    find_note_by_type,
    format_date_russian,
    parse_frontmatter,
    replace_vars_in_paragraph,
    resolve_wikilink,
)

logger = logging.getLogger(__name__)

NESTED_VAR_INNER = re.compile(r"\{\{VAR\|([^{|}]+)\|([^{|}]+)\}\}")


def resolve_nested_vars(text: str, variables: dict[str, str], max_depth: int = 5) -> str:
    """Рекурсивно разрешает вложенные {{VAR|ID|default}} изнутри наружу.

    Args:
        text: Текст с переменными.
        variables: Словарь переменных.
        max_depth: Максимальная глубина рекурсии.

    Returns:
        Текст с разрешёнными переменными.
    """
    for _ in range(max_depth):
        match = NESTED_VAR_INNER.search(text)
        if not match:
            break
        var_id = match.group(1).strip()
        default = match.group(2).strip()
        value = variables.get(var_id, default)
        text = text[:match.start()] + value + text[match.end():]
    return text


def replace_vars_in_paragraph_nested(
    paragraph: Any, variables: dict[str, str]
) -> int:
    """Заменяет переменные с поддержкой вложенности.

    Args:
        paragraph: Параграф Word.
        variables: Словарь переменных.

    Returns:
        Количество замен.
    """
    if not paragraph.runs:
        return 0

    full_text = "".join(run.text for run in paragraph.runs)
    if "{{VAR|" not in full_text:
        return 0

    # Сначала разрешаем вложенные
    resolved = resolve_nested_vars(full_text, variables)

    # Затем заменяем оставшиеся простые
    count = 0
    for match in VAR_PATTERN.finditer(resolved):
        var_id = match.group(1).strip()
        default_val = match.group(2).strip()
        value = variables.get(var_id, default_val)
        resolved = resolved.replace(match.group(0), value, 1)
        count += 1

    if resolved != full_text or count > 0:
        paragraph.runs[0].text = resolved
        for run in paragraph.runs[1:]:
            run.text = ""
        return max(count, 1)

    return 0


def format_number_with_spaces(num: int | float) -> str:
    """Форматирует число с разделителем тысяч пробелом.

    Args:
        num: Число для форматирования.

    Returns:
        Отформатированная строка (например, "150 000").
    """
    if isinstance(num, float) and num == int(num):
        num = int(num)
    if isinstance(num, int):
        return f"{num:,}".replace(",", " ")
    return f"{num:,.2f}".replace(",", " ")


def clone_row(table: Table, row_idx: int) -> _Row:
    """Клонирует строку таблицы Word.

    Args:
        table: Таблица Word.
        row_idx: Индекс строки для клонирования.

    Returns:
        Новая строка-клон.
    """
    source_row = table.rows[row_idx]
    new_row_el = copy.deepcopy(source_row._tr)
    source_row._tr.addnext(new_row_el)
    return _Row(new_row_el, table)


def find_service_table(doc: Any) -> Table | None:
    """Находит таблицу услуг по заголовку.

    Args:
        doc: Документ Word.

    Returns:
        Таблица или None.
    """
    for table in doc.tables:
        for row in table.rows[:2]:
            for cell in row.cells:
                text = cell.text.strip().lower()
                if "наименование" in text and "работ" in text:
                    return table
    return None


def process_service_table(
    table: Table,
    services: list[dict[str, Any]],
    variables: dict[str, str],
) -> None:
    """Обрабатывает динамическую таблицу услуг.

    Args:
        table: Таблица Word.
        services: Список услуг из YAML.
        variables: Словарь переменных.
    """
    if len(table.rows) < 2:
        return

    # Определяем строку-шаблон (первая строка данных после заголовка)
    header_rows = 1
    template_row_idx = header_rows

    # Определяем строки итогов (ищем "Итого" и "НДС")
    total_row_idx = None
    nds_row_idx = None
    for i, row in enumerate(table.rows):
        cell_text = row.cells[0].text.strip().lower() if row.cells else ""
        if "итого" in cell_text:
            total_row_idx = i
        if "ндс" in cell_text:
            nds_row_idx = i

    # Клонируем строки для услуг
    num_services = len(services)
    if num_services > 1:
        for _ in range(num_services - 1):
            clone_row(table, template_row_idx)

    # Заполняем строки услуг
    total_cost = 0
    for i, service in enumerate(services):
        row_idx = template_row_idx + i
        if row_idx >= len(table.rows):
            break

        svc_num = i + 1
        svc_vars = dict(variables)
        svc_vars["Номер варианта выполняемых работ"] = str(svc_num)
        svc_vars[f"Вариант выполняемых работ {svc_num}"] = str(service.get("название", ""))
        svc_vars[f"Вид отчётности варианта выполняемых работ {svc_num}"] = str(service.get("отчётность", ""))
        cost = service.get("стоимость", 0)
        total_cost += int(cost) if cost else 0
        svc_vars[f"Стоимость варианта выполняемых работ {svc_num}"] = format_number_with_spaces(cost)

        row = table.rows[row_idx]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_vars_in_paragraph_nested(paragraph, svc_vars)

    # Удаляем лишние строки шаблона (если услуг меньше чем строк)
    # (пропускаем — строки уже были клонированы по количеству)

    # Обновляем строку "Итого"
    if total_row_idx is not None:
        adj_total_idx = total_row_idx + (num_services - 1) if num_services > 1 else total_row_idx
        if adj_total_idx < len(table.rows):
            total_row = table.rows[adj_total_idx]
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    if "{{VAR|" in text or any(c.isdigit() for c in text):
                        # Заменяем сумму
                        total_vars = {"Итого": format_number_with_spaces(total_cost)}
                        replace_vars_in_paragraph_nested(paragraph, total_vars)

    # Обновляем строку "НДС 20%"
    if nds_row_idx is not None:
        adj_nds_idx = nds_row_idx + (num_services - 1) if num_services > 1 else nds_row_idx
        if adj_nds_idx < len(table.rows):
            nds = round(total_cost * 20 / 120)
            nds_row = table.rows[adj_nds_idx]
            for cell in nds_row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    if "{{VAR|" in text or "ндс" in text.lower():
                        nds_vars = {"НДС": format_number_with_spaces(nds)}
                        replace_vars_in_paragraph_nested(paragraph, nds_vars)


def generate_specification(
    vault: Path,
    contract_name: str,
    spec_number: int,
    template_path: Path | None = None,
    output_path: Path | None = None,
    dry_run: bool = False,
    check_grammar: bool = False,
) -> dict[str, Any]:
    """Генерирует .docx спецификацию.

    Args:
        vault: Путь к vault.
        contract_name: Имя заметки договора.
        spec_number: Номер спецификации (1-based).
        template_path: Путь к шаблону .docx.
        output_path: Путь для сохранения.
        dry_run: Только показать подстановки.
        check_grammar: Проверить грамматику.

    Returns:
        Результат генерации.
    """
    # 1. Чтение данных
    contract_file = resolve_wikilink(vault, contract_name)
    if contract_file is None:
        logger.error("Заметка договора не найдена: %s", contract_name)
        return {"success": False, "error": f"Заметка не найдена: {contract_name}"}

    contract_fm = parse_frontmatter(contract_file)
    if not contract_fm:
        return {"success": False, "error": "Пустой frontmatter"}

    counterparty_link = str(contract_fm.get("контрагент", ""))
    counterparty_file = resolve_wikilink(vault, counterparty_link)
    counterparty_fm = parse_frontmatter(counterparty_file) if counterparty_file else {}

    company_file = find_note_by_type(vault, "наша_компания")
    company_fm = parse_frontmatter(company_file) if company_file else {}

    # 2. Поиск спецификации
    specs = contract_fm.get("спецификации", [])
    if not specs or spec_number < 1 or spec_number > len(specs):
        return {"success": False, "error": f"Спецификация №{spec_number} не найдена"}

    spec = specs[spec_number - 1]

    # 3. Сборка переменных
    variables = build_contract_variables(contract_fm, company_fm, counterparty_fm)

    # Добавляем переменные спецификации
    spec_date = spec.get("дата", "")
    variables["Дата спецификации"] = format_date_russian(str(spec_date)) if spec_date else "«__» ________ 20__ г."

    # Доп. поля
    for field in ("место", "порядок_оплаты", "сроки", "гарантия", "доп_условия", "доп_информация"):
        if field in spec:
            variables[field] = str(spec[field])

    # Переменные для каждой услуги
    services = spec.get("услуги", [])
    total_cost = 0
    for i, svc in enumerate(services):
        svc_num = i + 1
        variables["Номер варианта выполняемых работ"] = str(svc_num)
        variables[f"Вариант выполняемых работ {svc_num}"] = str(svc.get("название", ""))
        variables[f"Вид отчётности варианта выполняемых работ {svc_num}"] = str(svc.get("отчётность", ""))
        cost = svc.get("стоимость", 0)
        total_cost += int(cost) if cost else 0
        variables[f"Стоимость варианта выполняемых работ {svc_num}"] = format_number_with_spaces(cost)

    variables["Итого"] = format_number_with_spaces(total_cost)
    variables["НДС"] = format_number_with_spaces(round(total_cost * 20 / 120))

    if dry_run:
        logger.info("=== DRY RUN: Таблица подстановок спецификации №%d ===", spec_number)
        result_table: list[dict[str, str]] = []
        for var_id, value in sorted(variables.items()):
            result_table.append({"variable": var_id, "value": value})
            logger.info("  %s → %s", var_id, value)
        return {
            "success": True,
            "dry_run": True,
            "spec_number": spec_number,
            "services_count": len(services),
            "total": format_number_with_spaces(total_cost),
            "nds": format_number_with_spaces(round(total_cost * 20 / 120)),
            "variables": result_table,
        }

    # 4. Подстановка
    if template_path is None:
        skill_root = Path(__file__).resolve().parent.parent
        template_path = skill_root / "assets" / "docx-templates" / "specification_services.docx"

    if not template_path.exists():
        return {"success": False, "error": f"Шаблон не найден: {template_path}"}

    doc = Document(str(template_path))

    # Обработка таблицы услуг
    svc_table = find_service_table(doc)
    if svc_table and services:
        process_service_table(svc_table, services, variables)

    # Обработка остальных параграфов с вложенными переменными
    total_replacements = 0
    for paragraph in doc.paragraphs:
        total_replacements += replace_vars_in_paragraph_nested(paragraph, variables)

    for table in doc.tables:
        if table is svc_table:
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_replacements += replace_vars_in_paragraph_nested(paragraph, variables)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    total_replacements += replace_vars_in_paragraph_nested(paragraph, variables)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    total_replacements += replace_vars_in_paragraph_nested(paragraph, variables)

    remaining = check_remaining_vars(doc)

    # 5. Сохранение
    contract_number = str(contract_fm.get("номер", "000")).replace("/", "-")
    if output_path is None:
        docs_dir = vault / "14-ВЛОЖЕНИЯ" / "Документы"
        docs_dir.mkdir(parents=True, exist_ok=True)
        output_path = docs_dir / f"Спецификация №{spec_number} к Договору {contract_number}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Спецификация сохранена: %s", output_path)

    result: dict[str, Any] = {
        "success": True,
        "output": str(output_path),
        "spec_number": spec_number,
        "services_count": len(services),
        "total": format_number_with_spaces(total_cost),
        "nds": format_number_with_spaces(round(total_cost * 20 / 120)),
        "replacements": total_replacements,
        "remaining_vars": remaining,
    }

    if check_grammar:
        try:
            from scripts.grammar_check import check_document
            warnings = check_document(output_path)
            result["grammar_warnings"] = warnings
        except ImportError:
            logger.warning("grammar_check.py недоступен")

    return result


def main() -> None:
    """Точка входа CLI."""
    parser = argparse.ArgumentParser(
        description="Генерация .docx спецификации к договору"
    )
    parser.add_argument("--vault", type=Path, required=True, help="Путь к vault")
    parser.add_argument("--contract", required=True, help="Имя заметки договора")
    parser.add_argument("--spec", type=int, required=True, help="Номер спецификации")
    parser.add_argument("--template", type=Path, default=None, help="Путь к шаблону .docx")
    parser.add_argument("--output", type=Path, default=None, help="Путь для сохранения")
    parser.add_argument("--dry-run", action="store_true", help="Только показать подстановки")
    parser.add_argument("--check-grammar", action="store_true", help="Проверить грамматику")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    vault = args.vault.resolve()
    if not vault.is_dir():
        logger.error("Vault не найден: %s", vault)
        sys.exit(1)

    result = generate_specification(
        vault=vault,
        contract_name=args.contract,
        spec_number=args.spec,
        template_path=args.template,
        output_path=args.output,
        dry_run=args.dry_run,
        check_grammar=args.check_grammar,
    )

    print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result.get("success") else 1)


if __name__ == "__main__":
    main()
