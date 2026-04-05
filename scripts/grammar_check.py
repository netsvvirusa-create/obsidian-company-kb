"""Проверка грамматики юридических документов .docx.

Проверяет 4 паттерна:
1. «в лице [должность] [ФИО]» → родительный падеж
2. «на основании [документ]» → родительный падеж
3. «именуемое/ый/ая» → согласование с ОПФ
4. «действующего/ей» → согласование с полом подписанта
"""

import argparse
import logging
import re
import sys
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

try:
    import pymorphy3

    MORPH = pymorphy3.MorphAnalyzer()
    HAS_PYMORPHY = True
except ImportError:
    MORPH = None  # type: ignore[assignment]
    HAS_PYMORPHY = False
    logger.warning("pymorphy3 не установлен — грамматическая проверка ограничена")

# Несклоняемые фамилии (распространённые окончания)
NON_DECLINABLE_ENDINGS = (
    "ко", "о", "е", "э", "и", "у", "ю", "а",
)


def is_genitive(word: str) -> bool:
    """Проверяет, стоит ли слово в родительном падеже.

    Args:
        word: Слово для проверки.

    Returns:
        True если слово в родительном падеже.
    """
    if not HAS_PYMORPHY or MORPH is None:
        return True
    parses = MORPH.parse(word)
    return any("gent" in str(p.tag) for p in parses)


def check_v_litse(text: str, para_num: int) -> list[dict[str, Any]]:
    """Проверяет паттерн «в лице [должность] [ФИО]».

    Args:
        text: Текст параграфа.
        para_num: Номер параграфа.

    Returns:
        Список предупреждений.
    """
    warnings: list[dict[str, Any]] = []
    if not HAS_PYMORPHY:
        return warnings

    pattern = re.compile(r"в\s+лице\s+(.+?)(?:,|\.|$)", re.IGNORECASE)
    for match in pattern.finditer(text):
        phrase = match.group(1).strip()
        words = phrase.split()
        for word in words[:5]:
            clean = re.sub(r"[,.]", "", word)
            if not clean or len(clean) < 2:
                continue
            if not is_genitive(clean):
                warnings.append({
                    "paragraph": para_num,
                    "pattern": "в лице",
                    "text": match.group(0).strip(),
                    "word": clean,
                    "suggestion": f"Слово «{clean}» может быть не в родительном падеже",
                    "severity": "warning",
                })
                break
    return warnings


def check_na_osnovanii(text: str, para_num: int) -> list[dict[str, Any]]:
    """Проверяет паттерн «на основании [документ]».

    Args:
        text: Текст параграфа.
        para_num: Номер параграфа.

    Returns:
        Список предупреждений.
    """
    warnings: list[dict[str, Any]] = []
    if not HAS_PYMORPHY:
        return warnings

    pattern = re.compile(r"на\s+основании\s+(\S+)", re.IGNORECASE)
    for match in pattern.finditer(text):
        word = match.group(1).strip().rstrip(",.")
        if len(word) < 2:
            continue
        if not is_genitive(word):
            warnings.append({
                "paragraph": para_num,
                "pattern": "на основании",
                "text": match.group(0).strip(),
                "word": word,
                "suggestion": f"Слово «{word}» должно быть в родительном падеже",
                "severity": "warning",
            })
    return warnings


def check_imenuemoe(text: str, para_num: int) -> list[dict[str, Any]]:
    """Проверяет согласование «именуемое/ый/ая» с ОПФ.

    ООО/АО/ПАО → «именуемое» (средний род)
    ИП → «именуемый» (мужской род)

    Args:
        text: Текст параграфа.
        para_num: Номер параграфа.

    Returns:
        Список предупреждений.
    """
    warnings: list[dict[str, Any]] = []

    neuter_opf = ("ООО", "АО", "ПАО", "НАО", "ЗАО", "ОАО")
    masc_opf = ("ИП",)

    pattern = re.compile(
        r"(ООО|АО|ПАО|НАО|ЗАО|ОАО|ИП)\s+.*?(именуем\w+)",
        re.IGNORECASE,
    )
    for match in pattern.finditer(text):
        opf = match.group(1).upper()
        word = match.group(2).lower()

        if opf in neuter_opf and word != "именуемое":
            warnings.append({
                "paragraph": para_num,
                "pattern": "именуемое",
                "text": match.group(0).strip()[:80],
                "word": match.group(2),
                "suggestion": f"Для {opf} следует использовать «именуемое» (средний род)",
                "severity": "warning",
            })
        elif opf in masc_opf and word != "именуемый":
            warnings.append({
                "paragraph": para_num,
                "pattern": "именуемый",
                "text": match.group(0).strip()[:80],
                "word": match.group(2),
                "suggestion": f"Для {opf} следует использовать «именуемый» (мужской род)",
                "severity": "warning",
            })

    return warnings


def check_deystvuyushego(text: str, para_num: int) -> list[dict[str, Any]]:
    """Проверяет согласование «действующего/ей» с полом подписанта.

    Определение по отчеству: -вна/-вны → «действующей», -вич → «действующего».

    Args:
        text: Текст параграфа.
        para_num: Номер параграфа.

    Returns:
        Список предупреждений.
    """
    warnings: list[dict[str, Any]] = []

    pattern = re.compile(
        r"(действующ\w+)\s+на\s+основании",
        re.IGNORECASE,
    )
    for match in pattern.finditer(text):
        deystvuyushiy = match.group(1).lower()

        # Ищем отчество в контексте (ФИО обычно перед «действующ*»)
        preceding = text[:match.start()]
        # Ищем слово, похожее на отчество
        fio_pattern = re.compile(r"(\S+вны|\S+вича|\S+вной|\S+вна|\S+вич)\b", re.IGNORECASE)
        otchestvo_match = fio_pattern.search(preceding)

        if otchestvo_match:
            otchestvo = otchestvo_match.group(1).lower()
            is_female = otchestvo.endswith(("вна", "вны", "вной"))
            is_male = otchestvo.endswith(("вич", "вича"))

            if is_female and deystvuyushiy != "действующей":
                warnings.append({
                    "paragraph": para_num,
                    "pattern": "действующего/ей",
                    "text": match.group(0).strip(),
                    "word": match.group(1),
                    "suggestion": "Для женского пола следует использовать «действующей»",
                    "severity": "warning",
                })
            elif is_male and deystvuyushiy != "действующего":
                warnings.append({
                    "paragraph": para_num,
                    "pattern": "действующего/ей",
                    "text": match.group(0).strip(),
                    "word": match.group(1),
                    "suggestion": "Для мужского пола следует использовать «действующего»",
                    "severity": "warning",
                })

    return warnings


def check_document(filepath: Path, verbose: bool = False) -> list[dict[str, Any]]:
    """Проверяет .docx документ на грамматические ошибки.

    Args:
        filepath: Путь к .docx файлу.
        verbose: Подробный вывод.

    Returns:
        Список предупреждений.
    """
    from docx import Document

    doc = Document(str(filepath))
    all_warnings: list[dict[str, Any]] = []

    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text
        if not text.strip():
            continue

        all_warnings.extend(check_v_litse(text, i))
        all_warnings.extend(check_na_osnovanii(text, i))
        all_warnings.extend(check_imenuemoe(text, i))
        all_warnings.extend(check_deystvuyushego(text, i))

    # Проверяем таблицы
    para_offset = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    para_offset += 1
                    text = paragraph.text
                    if not text.strip():
                        continue
                    all_warnings.extend(check_v_litse(text, para_offset))
                    all_warnings.extend(check_na_osnovanii(text, para_offset))
                    all_warnings.extend(check_imenuemoe(text, para_offset))
                    all_warnings.extend(check_deystvuyushego(text, para_offset))

    if verbose:
        for w in all_warnings:
            logger.info(
                "Параграф %d [%s]: %s — %s",
                w["paragraph"], w["pattern"], w["word"], w["suggestion"],
            )

    return all_warnings


def main() -> None:
    """Точка входа CLI."""
    parser = argparse.ArgumentParser(
        description="Проверка грамматики юридических документов .docx"
    )
    parser.add_argument("--file", type=Path, required=True, help="Путь к .docx файлу")
    parser.add_argument("--verbose", action="store_true", help="Подробный вывод")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    filepath = args.file.resolve()
    if not filepath.exists():
        logger.error("Файл не найден: %s", filepath)
        sys.exit(1)

    warnings = check_document(filepath, verbose=args.verbose)

    if warnings:
        print(f"\nНайдено предупреждений: {len(warnings)}")
        for w in warnings:
            print(f"  Параграф {w['paragraph']}: [{w['pattern']}] {w['suggestion']}")
    else:
        print("Грамматических проблем не найдено.")

    sys.exit(0)


if __name__ == "__main__":
    main()
