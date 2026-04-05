"""Генерация .docx договора из шаблона и данных Obsidian vault.

Читает заметку договора, разрешает [[wikilinks]] на контрагента и
заметку type:наша_компания, собирает словарь переменных и подставляет
в .docx шаблон с обработкой Word runs.
"""

import argparse
import copy
import json
import logging
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document

logger = logging.getLogger(__name__)

MONTH_NAMES_GENITIVE: dict[int, str] = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

VAR_PATTERN = re.compile(r"\{\{VAR\|([^|]+)\|([^}]+)\}\}")


def parse_frontmatter(filepath: Path) -> dict[str, Any]:
    """Парсит YAML frontmatter из .md файла.

    Args:
        filepath: Путь к файлу заметки.

    Returns:
        Словарь YAML frontmatter.
    """
    text = filepath.read_text(encoding="utf-8")
    parts = text.split("---", 2)
    if len(parts) < 3:
        return {}
    try:
        return yaml.safe_load(parts[1]) or {}
    except yaml.YAMLError:
        logger.warning("Ошибка парсинга YAML в %s", filepath)
        return {}


def resolve_wikilink(vault: Path, link: str) -> Path | None:
    """Находит файл по wikilink имени.

    Args:
        vault: Путь к vault.
        link: Имя из [[wikilink]].

    Returns:
        Путь к найденному файлу или None.
    """
    clean = link.strip().strip("[]").split("|")[0].split("#")[0].strip()
    if not clean:
        return None
    for md_file in vault.rglob("*.md"):
        if md_file.stem == clean:
            return md_file
    return None


def find_note_by_type(vault: Path, note_type: str) -> Path | None:
    """Находит заметку по значению поля type.

    Args:
        vault: Путь к vault.
        note_type: Значение поля type для поиска.

    Returns:
        Путь к найденному файлу или None.
    """
    for md_file in vault.rglob("*.md"):
        fm = parse_frontmatter(md_file)
        if fm.get("type") == note_type:
            return md_file
    return None


def format_date_russian(date_str: str) -> str:
    """Конвертирует YYYY-MM-DD в «DD» месяц YYYY г.

    Args:
        date_str: Дата в формате YYYY-MM-DD.

    Returns:
        Дата в формате «DD» месяц YYYY г.
    """
    try:
        if isinstance(date_str, date):
            d = date_str
        else:
            d = datetime.strptime(str(date_str), "%Y-%m-%d").date()
        month_name = MONTH_NAMES_GENITIVE[d.month]
        return f"\u00ab{d.day:02d}\u00bb {month_name} {d.year} г."
    except (ValueError, KeyError):
        return str(date_str)


def build_contract_variables(
    contract_fm: dict[str, Any],
    company_fm: dict[str, Any],
    counterparty_fm: dict[str, Any],
) -> dict[str, str]:
    """Собирает словарь переменных для подстановки в шаблон договора.

    Args:
        contract_fm: Frontmatter заметки договора.
        company_fm: Frontmatter заметки type:наша_компания.
        counterparty_fm: Frontmatter заметки контрагента.

    Returns:
        Словарь {ID переменной: значение}.
    """
    variables: dict[str, str] = {}

    # Общие
    variables["Номер договора"] = str(contract_fm.get("номер", "001"))
    date_val = contract_fm.get("дата_подписания", "")
    variables["Дата договора"] = format_date_russian(str(date_val)) if date_val else "«__» ________ 20__ г."

    # Исполнитель (наша компания)
    variables["Полный тип ОПФ лица Исполнителя"] = str(company_fm.get("опф_полная", "Общество с ограниченной ответственностью"))
    variables["Полное наименование ЮЛ Исполнителя"] = str(company_fm.get("название_полное", "«Наименование»"))
    variables["Краткий тип ОПФ лица Исполнителя"] = str(company_fm.get("опф_краткая", "ООО"))
    variables["Короткое наименование лица Исполнителя"] = str(company_fm.get("название_краткое", "«Наименование»"))
    variables["Вид стороны по договору 1"] = str(contract_fm.get("вид_стороны_1", "«Исполнитель»"))
    variables["Должность уполномоченного лица Исполнителя"] = str(contract_fm.get("должность_подписанта_нашего_рп", "Генерального директора"))
    variables["ФИО уполномоченного лица Исполнителя"] = str(contract_fm.get("фио_подписанта_нашего_рп", "Фамилии Имени Отчества"))
    variables["Основание действий уполномоченного лица Исполнителя"] = str(contract_fm.get("основание_подписанта_нашего_рп", "Устава"))
    variables["Физический адрес Исполнителя"] = str(company_fm.get("адрес", "Адрес"))
    variables["ИНН Исполнителя"] = str(company_fm.get("инн", "0000000000"))
    variables["КПП Исполнителя"] = str(company_fm.get("кпп", "000000000"))
    variables["Расчётный счёт Исполнителя"] = str(company_fm.get("расчётный_счёт", "00000000000000000000"))
    variables["Наименование банка Исполнителя"] = str(company_fm.get("банк", "Наименование банка"))
    variables["Пример БИК банка Исполнителя"] = str(company_fm.get("бик", "000000000"))
    variables["Пример Корр счёта банка Исполнителя"] = str(company_fm.get("корр_счёт", "00000000000000000000"))

    # Заказчик (контрагент)
    variables["Полный тип ОПФ лица Заказчика"] = str(counterparty_fm.get("опф_полная", "Общество с ограниченной ответственностью"))
    variables["Полное наименование ЮЛ Заказчика"] = str(counterparty_fm.get("название_полное", "«Наименование»"))
    variables["Краткий тип ОПФ лица Заказчика"] = str(counterparty_fm.get("опф_краткая", "ООО"))
    variables["Короткое наименование лица Заказчика"] = str(counterparty_fm.get("название_краткое", "«Наименование»"))
    variables["Вид стороны по договору 2"] = str(contract_fm.get("вид_стороны_2", "«Заказчик»"))
    variables["Должность уполномоченного лица Заказчика"] = str(contract_fm.get("должность_подписанта_контрагента_рп", "Генерального директора"))
    variables["ФИО уполномоченного лица Заказчика"] = str(contract_fm.get("фио_подписанта_контрагента_рп", "Фамилии Имени Отчества"))
    variables["Основание действий уполномоченного лица Заказчика"] = str(contract_fm.get("основание_подписанта_контрагента_рп", "Устава"))
    variables["Физический адрес Заказчика"] = str(counterparty_fm.get("фактический_адрес", counterparty_fm.get("юридический_адрес", "Адрес")))
    variables["ИНН Заказчика"] = str(counterparty_fm.get("инн", "0000000000"))
    variables["КПП Заказчика"] = str(counterparty_fm.get("кпп", "000000000"))
    variables["Расчётный счёт Заказчика"] = str(counterparty_fm.get("расчётный_счёт", "00000000000000000000"))
    variables["Наименование банка Заказчика"] = str(counterparty_fm.get("банк", "Наименование банка"))
    variables["Пример БИК банка Заказчика"] = str(counterparty_fm.get("бик", "000000000"))
    variables["Пример Корр счёта банка Заказчика"] = str(counterparty_fm.get("корр_счёт", "00000000000000000000"))

    return variables


def replace_vars_in_paragraph(paragraph: Any, variables: dict[str, str]) -> int:
    """Заменяет {{VAR|ID|default}} в параграфе Word.

    Word разбивает текст на runs. Переменная может быть разбита на
    несколько runs. Собираем полный текст, заменяем, записываем обратно.

    Args:
        paragraph: Параграф документа Word.
        variables: Словарь переменных для подстановки.

    Returns:
        Количество выполненных замен.
    """
    if not paragraph.runs:
        return 0

    full_text = "".join(run.text for run in paragraph.runs)
    if "{{VAR|" not in full_text:
        return 0

    count = 0
    for match in VAR_PATTERN.finditer(full_text):
        var_id = match.group(1).strip()
        default_val = match.group(2).strip()
        value = variables.get(var_id, default_val)
        full_text = full_text.replace(match.group(0), value, 1)
        count += 1

    if count > 0:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""

    return count


def process_document(doc: Any, variables: dict[str, str]) -> int:
    """Обрабатывает весь документ Word, заменяя переменные.

    Args:
        doc: Документ Word (python-docx Document).
        variables: Словарь переменных для подстановки.

    Returns:
        Общее количество замен.
    """
    total = 0

    # Параграфы
    for paragraph in doc.paragraphs:
        total += replace_vars_in_paragraph(paragraph, variables)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += replace_vars_in_paragraph(paragraph, variables)

    # Колонтитулы
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    total += replace_vars_in_paragraph(paragraph, variables)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    total += replace_vars_in_paragraph(paragraph, variables)

    return total


def check_remaining_vars(doc: Any) -> list[str]:
    """Проверяет наличие незаменённых переменных в документе.

    Args:
        doc: Документ Word.

    Returns:
        Список незаменённых переменных.
    """
    remaining: list[str] = []

    for paragraph in doc.paragraphs:
        text = "".join(run.text for run in paragraph.runs)
        for match in VAR_PATTERN.finditer(text):
            remaining.append(match.group(0))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    for match in VAR_PATTERN.finditer(text):
                        remaining.append(match.group(0))

    return remaining


def generate_contract(
    vault: Path,
    contract_name: str,
    template_path: Path | None = None,
    output_path: Path | None = None,
    dry_run: bool = False,
    check_grammar: bool = False,
) -> dict[str, Any]:
    """Генерирует .docx договор.

    Args:
        vault: Путь к vault.
        contract_name: Имя заметки договора.
        template_path: Путь к шаблону .docx (опционально).
        output_path: Путь для сохранения (опционально).
        dry_run: Только показать подстановки.
        check_grammar: Запустить проверку грамматики.

    Returns:
        Результат генерации в виде словаря.
    """
    # 1. Чтение заметки договора
    contract_file = resolve_wikilink(vault, contract_name)
    if contract_file is None:
        logger.error("Заметка договора не найдена: %s", contract_name)
        return {"success": False, "error": f"Заметка не найдена: {contract_name}"}

    contract_fm = parse_frontmatter(contract_file)
    if not contract_fm:
        logger.error("Пустой frontmatter в %s", contract_file)
        return {"success": False, "error": "Пустой frontmatter"}

    # 2. Разрешение wikilinks
    counterparty_link = str(contract_fm.get("контрагент", ""))
    counterparty_file = resolve_wikilink(vault, counterparty_link)
    counterparty_fm = parse_frontmatter(counterparty_file) if counterparty_file else {}

    company_file = find_note_by_type(vault, "наша_компания")
    company_fm = parse_frontmatter(company_file) if company_file else {}

    # 3. Сборка словаря
    variables = build_contract_variables(contract_fm, company_fm, counterparty_fm)

    if dry_run:
        logger.info("=== DRY RUN: Таблица подстановок ===")
        result_table: list[dict[str, str]] = []
        for var_id, value in sorted(variables.items()):
            result_table.append({"variable": var_id, "value": value})
            logger.info("  %s → %s", var_id, value)
        return {"success": True, "dry_run": True, "variables": result_table}

    # 4. Подстановка в .docx
    if template_path is None:
        skill_root = Path(__file__).resolve().parent.parent
        template_path = skill_root / "assets" / "docx-templates" / "contract_template.docx"

    if not template_path.exists():
        logger.error("Шаблон не найден: %s", template_path)
        return {"success": False, "error": f"Шаблон не найден: {template_path}"}

    doc = Document(str(template_path))
    total_replacements = process_document(doc, variables)

    remaining = check_remaining_vars(doc)
    if remaining:
        logger.warning("Незаменённые переменные: %s", remaining)

    # 5. Сохранение
    contract_number = str(contract_fm.get("номер", "000")).replace("/", "-")
    if output_path is None:
        docs_dir = vault / "14-ВЛОЖЕНИЯ" / "Документы"
        docs_dir.mkdir(parents=True, exist_ok=True)
        output_path = docs_dir / f"Договор {contract_number}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Документ сохранён: %s", output_path)

    result: dict[str, Any] = {
        "success": True,
        "output": str(output_path),
        "replacements": total_replacements,
        "remaining_vars": remaining,
    }

    # 6. Grammar check
    if check_grammar:
        try:
            from scripts.grammar_check import check_document
            warnings = check_document(output_path)
            result["grammar_warnings"] = warnings
        except ImportError:
            logger.warning("grammar_check.py недоступен")

    return result


def main() -> None:
    """Точка входа CLI."""
    parser = argparse.ArgumentParser(
        description="Генерация .docx договора из данных Obsidian vault"
    )
    parser.add_argument("--vault", type=Path, required=True, help="Путь к vault")
    parser.add_argument("--contract", required=True, help="Имя заметки договора")
    parser.add_argument("--template", type=Path, default=None, help="Путь к шаблону .docx")
    parser.add_argument("--output", type=Path, default=None, help="Путь для сохранения")
    parser.add_argument("--dry-run", action="store_true", help="Только показать подстановки")
    parser.add_argument("--check-grammar", action="store_true", help="Проверить грамматику")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    vault = args.vault.resolve()
    if not vault.is_dir():
        logger.error("Vault не найден: %s", vault)
        sys.exit(1)

    result = generate_contract(
        vault=vault,
        contract_name=args.contract,
        template_path=args.template,
        output_path=args.output,
        dry_run=args.dry_run,
        check_grammar=args.check_grammar,
    )

    print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result.get("success") else 1)


if __name__ == "__main__":
    main()
