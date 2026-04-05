"""Tests for scripts.generate_specification module."""

from pathlib import Path

import pytest

from scripts.generate_contract import (
    build_contract_variables,
    parse_frontmatter,
)
from scripts.generate_specification import (
    format_number_with_spaces,
    generate_specification,
    replace_vars_in_paragraph_nested,
    resolve_nested_vars,
)

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.mark.unit
class TestResolveNestedVars:
    def test_inner_vars_resolved_first(self):
        text = (
            "{{VAR|\u0412\u0430\u0440\u0438\u0430\u043d\u0442 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u0435\u043c\u044b\u0445 \u0440\u0430\u0431\u043e\u0442 "
            "{{VAR|\u041d\u043e\u043c\u0435\u0440 \u0432\u0430\u0440\u0438\u0430\u043d\u0442\u0430 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u0435\u043c\u044b\u0445 \u0440\u0430\u0431\u043e\u0442|1}}"
            "|default}}"
        )
        variables = {
            "\u041d\u043e\u043c\u0435\u0440 \u0432\u0430\u0440\u0438\u0430\u043d\u0442\u0430 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u0435\u043c\u044b\u0445 \u0440\u0430\u0431\u043e\u0442": "1",
            "\u0412\u0430\u0440\u0438\u0430\u043d\u0442 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u0435\u043c\u044b\u0445 \u0440\u0430\u0431\u043e\u0442 1": "\u0413\u043e\u0434\u043e\u0432\u0430\u044f \u043f\u043e\u0434\u0434\u0435\u0440\u0436\u043a\u0430",
        }
        result = resolve_nested_vars(text, variables)
        assert result == "\u0413\u043e\u0434\u043e\u0432\u0430\u044f \u043f\u043e\u0434\u0434\u0435\u0440\u0436\u043a\u0430"


@pytest.mark.unit
class TestFormatNumberWithSpaces:
    def test_150000(self):
        assert format_number_with_spaces(150000) == "150 000"

    def test_1500000(self):
        assert format_number_with_spaces(1500000) == "1 500 000"

    def test_small_number(self):
        assert format_number_with_spaces(100) == "100"

    def test_float_whole(self):
        assert format_number_with_spaces(150000.0) == "150 000"


@pytest.mark.unit
class TestGenerateSpecificationDryRun:
    def test_dry_run_returns_variables(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="\u0414\u043e\u0433\u043e\u0432\u043e\u0440 \u2116001",
            spec_number=1,
            dry_run=True,
        )

        assert result["success"] is True
        assert result["dry_run"] is True
        assert result["spec_number"] == 1
        assert result["services_count"] == 1
        assert isinstance(result["variables"], list)
        assert len(result["variables"]) > 0


@pytest.mark.unit
class TestOneService:
    def test_one_service_in_spec(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="\u0414\u043e\u0433\u043e\u0432\u043e\u0440 \u2116001",
            spec_number=1,
            dry_run=True,
        )

        assert result["success"] is True
        assert result["services_count"] == 1
        assert result["total"] == "150 000"


@pytest.mark.unit
class TestFiveServices:
    def test_five_services_total_and_nds(self, tmp_vault, sample_company_note, sample_counterparty_note):
        note = tmp_vault / "02-\u0414\u041e\u0413\u041e\u0412\u041e\u0420\u042b" / "FiveServices.md"
        note.write_text(
            '---\n'
            'title: "FiveServices"\n'
            'type: \u0434\u043e\u0433\u043e\u0432\u043e\u0440\n'
            '\u043d\u043e\u043c\u0435\u0440: "FS-001"\n'
            '\u043a\u043e\u043d\u0442\u0440\u0430\u0433\u0435\u043d\u0442: "[[ООО Пример]]"\n'
            '\u0434\u0430\u0442\u0430_\u043f\u043e\u0434\u043f\u0438\u0441\u0430\u043d\u0438\u044f: 2026-04-05\n'
            '\u0441\u043f\u0435\u0446\u0438\u0444\u0438\u043a\u0430\u0446\u0438\u0438:\n'
            '  - \u043d\u043e\u043c\u0435\u0440: 1\n'
            '    \u0434\u0430\u0442\u0430: "2026-04-10"\n'
            '    \u0442\u0438\u043f: \u0443\u0441\u043b\u0443\u0433\u0438\n'
            '    \u0443\u0441\u043b\u0443\u0433\u0438:\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 1"\n'
            '        \u043e\u0442\u0447\u0451\u0442\u043d\u043e\u0441\u0442\u044c: "\u0410\u043a\u0442"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 60000\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 2"\n'
            '        \u043e\u0442\u0447\u0451\u0442\u043d\u043e\u0441\u0442\u044c: "\u0410\u043a\u0442"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 60000\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 3"\n'
            '        \u043e\u0442\u0447\u0451\u0442\u043d\u043e\u0441\u0442\u044c: "\u0410\u043a\u0442"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 60000\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 4"\n'
            '        \u043e\u0442\u0447\u0451\u0442\u043d\u043e\u0441\u0442\u044c: "\u0410\u043a\u0442"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 60000\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 5"\n'
            '        \u043e\u0442\u0447\u0451\u0442\u043d\u043e\u0441\u0442\u044c: "\u0410\u043a\u0442"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 60000\n'
            '    \u043c\u0435\u0441\u0442\u043e: "\u0433. \u041c\u043e\u0441\u043a\u0432\u0430"\n'
            '---\n\n# FiveServices\n',
            encoding="utf-8",
        )

        result = generate_specification(
            vault=tmp_vault,
            contract_name="FiveServices",
            spec_number=1,
            dry_run=True,
        )

        assert result["success"] is True
        assert result["services_count"] == 5
        assert result["total"] == "300 000"
        # NDS = 300000 * 20 / 120 = 50000
        assert result["nds"] == "50 000"


@pytest.mark.unit
class TestNdsCalculation:
    def test_nds_formula(self):
        # NDS = total_cost * 20 / 120
        total_cost = 300000
        nds = round(total_cost * 20 / 120)
        assert nds == 50000

        total_cost2 = 150000
        nds2 = round(total_cost2 * 20 / 120)
        assert nds2 == 25000


@pytest.mark.unit
class TestAutoincrementSpecNumber:
    def test_spec_number_parameter(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
    ):
        note = tmp_vault / "02-\u0414\u041e\u0413\u041e\u0412\u041e\u0420\u042b" / "MultiSpec.md"
        note.write_text(
            '---\n'
            'title: "MultiSpec"\n'
            'type: \u0434\u043e\u0433\u043e\u0432\u043e\u0440\n'
            '\u043d\u043e\u043c\u0435\u0440: "MS-001"\n'
            '\u043a\u043e\u043d\u0442\u0440\u0430\u0433\u0435\u043d\u0442: "[[ООО Пример]]"\n'
            '\u0434\u0430\u0442\u0430_\u043f\u043e\u0434\u043f\u0438\u0441\u0430\u043d\u0438\u044f: 2026-04-05\n'
            '\u0441\u043f\u0435\u0446\u0438\u0444\u0438\u043a\u0430\u0446\u0438\u0438:\n'
            '  - \u043d\u043e\u043c\u0435\u0440: 1\n'
            '    \u0434\u0430\u0442\u0430: "2026-04-10"\n'
            '    \u0442\u0438\u043f: \u0443\u0441\u043b\u0443\u0433\u0438\n'
            '    \u0443\u0441\u043b\u0443\u0433\u0438:\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 A"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 100000\n'
            '  - \u043d\u043e\u043c\u0435\u0440: 2\n'
            '    \u0434\u0430\u0442\u0430: "2026-05-10"\n'
            '    \u0442\u0438\u043f: \u0443\u0441\u043b\u0443\u0433\u0438\n'
            '    \u0443\u0441\u043b\u0443\u0433\u0438:\n'
            '      - \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u0435: "\u0423\u0441\u043b\u0443\u0433\u0430 B"\n'
            '        \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c: 200000\n'
            '---\n\n# MultiSpec\n',
            encoding="utf-8",
        )

        result1 = generate_specification(
            vault=tmp_vault, contract_name="MultiSpec", spec_number=1, dry_run=True,
        )
        assert result1["success"] is True
        assert result1["spec_number"] == 1
        assert result1["total"] == "100 000"

        result2 = generate_specification(
            vault=tmp_vault, contract_name="MultiSpec", spec_number=2, dry_run=True,
        )
        assert result2["success"] is True
        assert result2["spec_number"] == 2
        assert result2["total"] == "200 000"


@pytest.mark.unit
class TestAdditionalFields:
    def test_additional_fields_filled(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="\u0414\u043e\u0433\u043e\u0432\u043e\u0440 \u2116001",
            spec_number=1,
            dry_run=True,
        )

        assert result["success"] is True
        variables = {v["variable"]: v["value"] for v in result["variables"]}
        assert variables["\u043c\u0435\u0441\u0442\u043e"] == "\u0433. \u041c\u043e\u0441\u043a\u0432\u0430"
        assert variables["\u043f\u043e\u0440\u044f\u0434\u043e\u043a_\u043e\u043f\u043b\u0430\u0442\u044b"] == "100% \u043f\u0440\u0435\u0434\u043e\u043f\u043b\u0430\u0442\u0430"
        assert variables["\u0441\u0440\u043e\u043a\u0438"] == "1 \u043c\u0435\u0441\u044f\u0446 \u0441 \u043c\u043e\u043c\u0435\u043d\u0442\u0430 \u043e\u043f\u043b\u0430\u0442\u044b"
        assert variables["\u0433\u0430\u0440\u0430\u043d\u0442\u0438\u044f"] == "1 \u0433\u043e\u0434"


# -----------------------------------------------------------------------
# NEW TESTS: resolve_nested_vars edge cases
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestResolveNestedVarsEdgeCases:
    def test_no_vars_returns_unchanged(self):
        text = "Просто текст без переменных"
        result = resolve_nested_vars(text, {})
        assert result == text

    def test_unknown_var_uses_default(self):
        text = "{{VAR|Несуществующая|fallback}}"
        result = resolve_nested_vars(text, {})
        assert result == "fallback"

    def test_known_var_replaces_value(self):
        text = "{{VAR|Ключ|дефолт}}"
        result = resolve_nested_vars(text, {"Ключ": "Значение"})
        assert result == "Значение"

    def test_multiple_nested_vars(self):
        text = "{{VAR|A|da}} и {{VAR|B|db}}"
        result = resolve_nested_vars(text, {"A": "1", "B": "2"})
        assert result == "1 и 2"

    def test_max_depth_reached(self):
        """If nesting exceeds max_depth, inner vars may remain."""
        text = "{{VAR|A|x}}"
        result = resolve_nested_vars(text, {}, max_depth=0)
        # With max_depth=0, the loop body never executes
        assert "{{VAR|" in result


# -----------------------------------------------------------------------
# NEW TESTS: clone_row, find_service_table, process_service_table
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestCloneRow:
    def test_clone_row_adds_row(self, specification_template_path):
        from scripts.generate_specification import clone_row, find_service_table
        from docx import Document

        doc = Document(str(specification_template_path))
        table = find_service_table(doc)
        if table is None:
            pytest.skip("No service table found in template")

        original_count = len(table.rows)
        clone_row(table, 1)
        assert len(table.rows) == original_count + 1


@pytest.mark.unit
class TestFindServiceTable:
    def test_finds_table_in_template(self, specification_template_path):
        from scripts.generate_specification import find_service_table
        from docx import Document

        doc = Document(str(specification_template_path))
        table = find_service_table(doc)
        assert table is not None

    def test_returns_none_for_empty_doc(self):
        from scripts.generate_specification import find_service_table
        from docx import Document

        doc = Document()
        doc.add_paragraph("No tables here")
        assert find_service_table(doc) is None


@pytest.mark.unit
class TestProcessServiceTable:
    def test_fills_services_into_table(self, specification_template_path):
        from scripts.generate_specification import (
            find_service_table,
            process_service_table,
            format_number_with_spaces,
        )
        from docx import Document

        doc = Document(str(specification_template_path))
        table = find_service_table(doc)
        if table is None:
            pytest.skip("No service table found in template")

        services = [
            {"название": "Услуга А", "отчётность": "Акт", "стоимость": 100000},
            {"название": "Услуга Б", "отчётность": "Акт", "стоимость": 200000},
        ]
        variables = {
            "Номер варианта выполняемых работ": "1",
            "Вариант выполняемых работ 1": "Услуга А",
            "Стоимость варианта выполняемых работ 1": format_number_with_spaces(100000),
            "Вариант выполняемых работ 2": "Услуга Б",
            "Стоимость варианта выполняемых работ 2": format_number_with_spaces(200000),
            "Итого": format_number_with_spaces(300000),
            "НДС": format_number_with_spaces(50000),
        }

        rows_before = len(table.rows)
        process_service_table(table, services, variables)
        # Should have cloned one extra row for the second service
        assert len(table.rows) >= rows_before + 1

    def test_single_service_no_clone(self, specification_template_path):
        from scripts.generate_specification import (
            find_service_table,
            process_service_table,
            format_number_with_spaces,
        )
        from docx import Document

        doc = Document(str(specification_template_path))
        table = find_service_table(doc)
        if table is None:
            pytest.skip("No service table found in template")

        services = [
            {"название": "Услуга А", "отчётность": "Акт", "стоимость": 100000},
        ]
        variables = {
            "Номер варианта выполняемых работ": "1",
            "Вариант выполняемых работ 1": "Услуга А",
            "Стоимость варианта выполняемых работ 1": format_number_with_spaces(100000),
            "Итого": format_number_with_spaces(100000),
            "НДС": format_number_with_spaces(round(100000 * 20 / 120)),
        }

        rows_before = len(table.rows)
        process_service_table(table, services, variables)
        # Single service should not clone
        assert len(table.rows) == rows_before


# -----------------------------------------------------------------------
# NEW TESTS: replace_vars_in_paragraph_nested
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestReplaceVarsInParagraphNestedAdditional:
    def test_replaces_simple_var_in_paragraph(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("{{VAR|Ключ|по-умолчанию}}")
        count = replace_vars_in_paragraph_nested(para, {"Ключ": "Значение"})
        assert count >= 1
        assert para.runs[0].text == "Значение"

    def test_no_vars_returns_zero(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Просто текст")
        count = replace_vars_in_paragraph_nested(para, {})
        assert count == 0

    def test_empty_runs_returns_zero(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        # Clear all runs
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        count = replace_vars_in_paragraph_nested(para, {"A": "B"})
        assert count == 0


# -----------------------------------------------------------------------
# NEW TESTS: full generate_specification pipeline (not dry_run)
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestGenerateSpecificationFull:
    def test_generates_docx_file(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
        specification_template_path,
    ):
        output = tmp_vault / "14-ВЛОЖЕНИЯ" / "Документы" / "test_spec.docx"
        result = generate_specification(
            vault=tmp_vault,
            contract_name="Договор №001",
            spec_number=1,
            template_path=specification_template_path,
            output_path=output,
            dry_run=False,
        )
        assert result["success"] is True
        assert output.exists()
        assert result["replacements"] >= 0
        assert result["spec_number"] == 1
        assert result["services_count"] == 1
        assert "remaining_vars" in result

    def test_five_services_generates_docx(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        specification_template_path,
    ):
        note = tmp_vault / "02-ДОГОВОРЫ" / "FiveSvcGen.md"
        note.write_text(
            '---\n'
            'title: "FiveSvcGen"\n'
            'type: договор\n'
            'номер: "FSG-001"\n'
            'контрагент: "[[ООО Пример]]"\n'
            'дата_подписания: 2026-04-05\n'
            'спецификации:\n'
            '  - номер: 1\n'
            '    дата: "2026-04-10"\n'
            '    тип: услуги\n'
            '    услуги:\n'
            '      - название: "Услуга 1"\n'
            '        отчётность: "Акт"\n'
            '        стоимость: 50000\n'
            '      - название: "Услуга 2"\n'
            '        отчётность: "Акт"\n'
            '        стоимость: 50000\n'
            '      - название: "Услуга 3"\n'
            '        отчётность: "Акт"\n'
            '        стоимость: 50000\n'
            '      - название: "Услуга 4"\n'
            '        отчётность: "Акт"\n'
            '        стоимость: 50000\n'
            '      - название: "Услуга 5"\n'
            '        отчётность: "Акт"\n'
            '        стоимость: 50000\n'
            '    место: "г. Москва"\n'
            '---\n\n# FiveSvcGen\n',
            encoding="utf-8",
        )

        output = tmp_vault / "14-ВЛОЖЕНИЯ" / "Документы" / "test_5svc.docx"
        result = generate_specification(
            vault=tmp_vault,
            contract_name="FiveSvcGen",
            spec_number=1,
            template_path=specification_template_path,
            output_path=output,
            dry_run=False,
        )
        assert result["success"] is True
        assert output.exists()
        assert result["services_count"] == 5
        assert result["total"] == "250 000"

    def test_contract_not_found(self, tmp_vault, sample_company_note):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="НесуществующийДоговор",
            spec_number=1,
        )
        assert result["success"] is False
        assert "не найдена" in result["error"]

    def test_invalid_spec_number(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="Договор №001",
            spec_number=99,
        )
        assert result["success"] is False
        assert "не найдена" in result["error"]

    def test_spec_number_zero(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="Договор №001",
            spec_number=0,
        )
        assert result["success"] is False

    def test_missing_template_error(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
    ):
        result = generate_specification(
            vault=tmp_vault,
            contract_name="Договор №001",
            spec_number=1,
            template_path=Path("/nonexistent/template.docx"),
        )
        assert result["success"] is False
        assert "Шаблон не найден" in result["error"]

    def test_empty_frontmatter_error(self, tmp_vault, sample_company_note):
        bad_note = tmp_vault / "02-ДОГОВОРЫ" / "BadNote.md"
        bad_note.write_text("# No frontmatter\n", encoding="utf-8")
        result = generate_specification(
            vault=tmp_vault,
            contract_name="BadNote",
            spec_number=1,
        )
        assert result["success"] is False

    def test_default_output_path(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
        specification_template_path,
    ):
        """When output_path is None, the file is saved to default location."""
        result = generate_specification(
            vault=tmp_vault,
            contract_name="Договор №001",
            spec_number=1,
            template_path=specification_template_path,
            output_path=None,
        )
        assert result["success"] is True
        assert Path(result["output"]).exists()


# -----------------------------------------------------------------------
# NEW TESTS: format_number_with_spaces additional cases
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestFormatNumberWithSpacesAdditional:
    def test_zero(self):
        assert format_number_with_spaces(0) == "0"

    def test_float_fractional(self):
        result = format_number_with_spaces(1500.55)
        assert "1 500" in result
        assert "55" in result

    def test_negative(self):
        result = format_number_with_spaces(-1000)
        assert "-1 000" == result
