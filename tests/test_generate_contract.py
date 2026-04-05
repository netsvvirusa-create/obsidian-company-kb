"""Tests for scripts.generate_contract module."""

import re
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from scripts.generate_contract import (
    VAR_PATTERN,
    build_contract_variables,
    check_remaining_vars,
    find_note_by_type,
    format_date_russian,
    generate_contract,
    parse_frontmatter,
    process_document,
    replace_vars_in_paragraph,
    resolve_wikilink,
)

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.mark.unit
class TestFormatDateRussian:
    def test_format_date_russian_april(self):
        assert format_date_russian("2026-04-05") == "\u00ab05\u00bb \u0430\u043f\u0440\u0435\u043b\u044f 2026 \u0433."

    def test_format_date_russian_january(self):
        assert format_date_russian("2026-01-15") == "\u00ab15\u00bb \u044f\u043d\u0432\u0430\u0440\u044f 2026 \u0433."


@pytest.mark.unit
class TestParseFrontmatter:
    def test_parse_frontmatter_valid(self, sample_contract_note):
        fm = parse_frontmatter(sample_contract_note)
        assert isinstance(fm, dict)
        assert fm["type"] == "\u0434\u043e\u0433\u043e\u0432\u043e\u0440"
        assert fm["\u043d\u043e\u043c\u0435\u0440"] == "001"
        assert fm["\u0441\u0443\u043c\u043c\u0430"] == 150000

    def test_parse_frontmatter_no_frontmatter(self, tmp_path):
        note = tmp_path / "empty.md"
        note.write_text("No frontmatter here", encoding="utf-8")
        fm = parse_frontmatter(note)
        assert fm == {}


@pytest.mark.unit
class TestResolveWikilink:
    def test_resolve_existing(self, tmp_vault, sample_counterparty_note):
        result = resolve_wikilink(tmp_vault, "[[ООО Пример]]")
        assert result is not None
        assert result.stem == "ООО Пример"

    def test_resolve_missing(self, tmp_vault):
        result = resolve_wikilink(tmp_vault, "[[Несуществующая заметка]]")
        assert result is None


@pytest.mark.unit
class TestFindNoteByType:
    def test_find_company_note(self, tmp_vault, sample_company_note):
        result = find_note_by_type(tmp_vault, "\u043d\u0430\u0448\u0430_\u043a\u043e\u043c\u043f\u0430\u043d\u0438\u044f")
        assert result is not None
        assert result.name == "ООО «Цифратроника».md"

    def test_find_missing_type(self, tmp_vault):
        result = find_note_by_type(tmp_vault, "\u043d\u0435\u0441\u0443\u0449\u0435\u0441\u0442\u0432\u0443\u044e\u0449\u0438\u0439_\u0442\u0438\u043f")
        assert result is None


@pytest.mark.unit
class TestBuildContractVariables:
    def test_returns_32_keys(self, sample_contract_note, sample_company_note, sample_counterparty_note):
        contract_fm = parse_frontmatter(sample_contract_note)
        company_fm = parse_frontmatter(sample_company_note)
        counterparty_fm = parse_frontmatter(sample_counterparty_note)

        variables = build_contract_variables(contract_fm, company_fm, counterparty_fm)

        assert isinstance(variables, dict)
        assert len(variables) == 32
        assert variables["\u041d\u043e\u043c\u0435\u0440 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430"] == "001"
        assert "\u0430\u043f\u0440\u0435\u043b\u044f" in variables["\u0414\u0430\u0442\u0430 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430"]
        assert variables["\u0418\u041d\u041d \u0418\u0441\u043f\u043e\u043b\u043d\u0438\u0442\u0435\u043b\u044f"] == "5024251588"
        assert variables["\u0418\u041d\u041d \u0417\u0430\u043a\u0430\u0437\u0447\u0438\u043a\u0430"] == "7701234567"


@pytest.mark.unit
class TestGenerateContractDryRun:
    def test_dry_run_returns_table(
        self, tmp_vault, sample_company_note, sample_counterparty_note, sample_contract_note
    ):
        result = generate_contract(
            vault=tmp_vault,
            contract_name="\u0414\u043e\u0433\u043e\u0432\u043e\u0440 \u2116001",
            dry_run=True,
        )

        assert result["success"] is True
        assert result["dry_run"] is True
        assert isinstance(result["variables"], list)
        assert len(result["variables"]) > 0

        # No file should be created
        docs_dir = tmp_vault / "14-\u0412\u041b\u041e\u0416\u0415\u041d\u0418\u042f" / "\u0414\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u044b"
        docx_files = list(docs_dir.glob("*.docx"))
        assert len(docx_files) == 0


@pytest.mark.unit
@pytest.mark.skipif(
    not (FIXTURES_DIR / "contract_template.docx").exists(),
    reason="contract_template.docx not found in fixtures",
)
class TestGenerateContractFull:
    def test_generates_docx(
        self,
        tmp_vault,
        sample_company_note,
        sample_counterparty_note,
        sample_contract_note,
        contract_template_path,
    ):
        output = tmp_vault / "output" / "test_contract.docx"
        result = generate_contract(
            vault=tmp_vault,
            contract_name="\u0414\u043e\u0433\u043e\u0432\u043e\u0440 \u2116001",
            template_path=contract_template_path,
            output_path=output,
        )

        assert result["success"] is True
        assert output.exists()

        # Check no remaining {{VAR}} in the generated document
        doc = Document(str(output))
        remaining = check_remaining_vars(doc)
        # remaining may contain vars not in our 32 keys, but core ones should be replaced
        for item in remaining:
            # None of the 32 build_contract_variables keys should remain
            match = VAR_PATTERN.search(item)
            if match:
                var_id = match.group(1).strip()
                # We only assert known keys are replaced; template may have extra vars
                assert var_id not in (
                    "\u041d\u043e\u043c\u0435\u0440 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430",
                    "\u0418\u041d\u041d \u0418\u0441\u043f\u043e\u043b\u043d\u0438\u0442\u0435\u043b\u044f",
                    "\u0418\u041d\u041d \u0417\u0430\u043a\u0430\u0437\u0447\u0438\u043a\u0430",
                ), f"Variable {var_id} was not replaced"


@pytest.mark.unit
class TestGenerateContractMissingData:
    def test_missing_data_uses_defaults(self, tmp_vault):
        # Create a minimal contract note with almost no data
        note = tmp_vault / "02-\u0414\u041e\u0413\u041e\u0412\u041e\u0420\u042b" / "Minimal.md"
        note.write_text(
            '---\ntitle: "Minimal"\ntype: \u0434\u043e\u0433\u043e\u0432\u043e\u0440\n---\n\n# Minimal\n',
            encoding="utf-8",
        )

        result = generate_contract(
            vault=tmp_vault,
            contract_name="Minimal",
            dry_run=True,
        )

        assert result["success"] is True
        variables = {v["variable"]: v["value"] for v in result["variables"]}
        # Defaults should be applied
        assert variables["\u041d\u043e\u043c\u0435\u0440 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430"] == "001"
        assert variables["\u041f\u043e\u043b\u043d\u044b\u0439 \u0442\u0438\u043f \u041e\u041f\u0424 \u043b\u0438\u0446\u0430 \u0417\u0430\u043a\u0430\u0437\u0447\u0438\u043a\u0430"] == "\u041e\u0431\u0449\u0435\u0441\u0442\u0432\u043e \u0441 \u043e\u0433\u0440\u0430\u043d\u0438\u0447\u0435\u043d\u043d\u043e\u0439 \u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0441\u0442\u044c\u044e"


@pytest.mark.unit
class TestFilenameSlashReplacement:
    def test_slash_in_contract_number(
        self, tmp_vault, sample_company_note, sample_counterparty_note
    ):
        note = tmp_vault / "02-\u0414\u041e\u0413\u041e\u0412\u041e\u0420\u042b" / "SlashTest.md"
        note.write_text(
            '---\n'
            'title: "SlashTest"\n'
            'type: \u0434\u043e\u0433\u043e\u0432\u043e\u0440\n'
            '\u043d\u043e\u043c\u0435\u0440: "001/2026"\n'
            '\u043a\u043e\u043d\u0442\u0440\u0430\u0433\u0435\u043d\u0442: "[[ООО Пример]]"\n'
            '\u0434\u0430\u0442\u0430_\u043f\u043e\u0434\u043f\u0438\u0441\u0430\u043d\u0438\u044f: 2026-04-05\n'
            '---\n\n# SlashTest\n',
            encoding="utf-8",
        )

        # Use a real template if available, otherwise skip the file check
        template = FIXTURES_DIR / "contract_template.docx"
        if not template.exists():
            pytest.skip("contract_template.docx not found")

        result = generate_contract(
            vault=tmp_vault,
            contract_name="SlashTest",
            template_path=template,
        )

        assert result["success"] is True
        output_path = Path(result["output"])
        assert "/" not in output_path.name
        assert "-" in output_path.name or "001" in output_path.name


@pytest.mark.unit
class TestSameIdAllReplaced:
    def test_all_occurrences_replaced(self):
        """If the same VAR ID appears 10 times, all 10 should be replaced."""
        doc = Document()
        var_text = "{{VAR|\u041d\u043e\u043c\u0435\u0440 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430|000}}"
        paragraph_text = " / ".join([var_text] * 10)
        para = doc.add_paragraph(paragraph_text)

        variables = {"\u041d\u043e\u043c\u0435\u0440 \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0430": "42"}
        count = replace_vars_in_paragraph(para, variables)

        assert count == 10
        result_text = "".join(run.text for run in para.runs)
        assert "{{VAR|" not in result_text
        assert result_text.count("42") == 10
