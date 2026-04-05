"""Tests for scripts.grammar_check module."""

import pytest

pymorphy3 = pytest.importorskip("pymorphy3")

from scripts.grammar_check import (
    check_deystvuyushego,
    check_imenuemoe,
    check_na_osnovanii,
    check_v_litse,
)


@pytest.mark.unit
class TestCheckVLitse:
    def test_correct_genitive(self):
        text = "\u0432 \u043b\u0438\u0446\u0435 \u0413\u0435\u043d\u0435\u0440\u0430\u043b\u044c\u043d\u043e\u0433\u043e \u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440\u0430"
        warnings = check_v_litse(text, 1)
        assert len(warnings) == 0

    def test_incorrect_nominative(self):
        text = "\u0432 \u043b\u0438\u0446\u0435 \u0413\u0435\u043d\u0435\u0440\u0430\u043b\u044c\u043d\u044b\u0439 \u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440"
        warnings = check_v_litse(text, 1)
        assert len(warnings) > 0
        assert warnings[0]["pattern"] == "\u0432 \u043b\u0438\u0446\u0435"
        assert warnings[0]["severity"] == "warning"


@pytest.mark.unit
class TestCheckNaOsnovanii:
    def test_correct_genitive(self):
        text = "\u043d\u0430 \u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0438 \u0423\u0441\u0442\u0430\u0432\u0430"
        warnings = check_na_osnovanii(text, 1)
        assert len(warnings) == 0

    def test_incorrect_nominative(self):
        text = "\u043d\u0430 \u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0438 \u0423\u0441\u0442\u0430\u0432"
        warnings = check_na_osnovanii(text, 1)
        assert len(warnings) > 0
        assert warnings[0]["pattern"] == "\u043d\u0430 \u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0438"


@pytest.mark.unit
class TestCheckImenuemoe:
    def test_correct_neuter_for_ooo(self):
        text = '\u041e\u041e\u041e \u00ab\u0426\u0438\u0444\u0440\u0430\u0442\u0440\u043e\u043d\u0438\u043a\u0430\u00bb, \u0438\u043c\u0435\u043d\u0443\u0435\u043c\u043e\u0435 \u0432 \u0434\u0430\u043b\u044c\u043d\u0435\u0439\u0448\u0435\u043c'
        warnings = check_imenuemoe(text, 1)
        assert len(warnings) == 0

    def test_incorrect_masculine_for_ooo(self):
        text = '\u041e\u041e\u041e \u00ab\u0426\u0438\u0444\u0440\u0430\u0442\u0440\u043e\u043d\u0438\u043a\u0430\u00bb, \u0438\u043c\u0435\u043d\u0443\u0435\u043c\u044b\u0439 \u0432 \u0434\u0430\u043b\u044c\u043d\u0435\u0439\u0448\u0435\u043c'
        warnings = check_imenuemoe(text, 1)
        assert len(warnings) > 0
        assert "\u0438\u043c\u0435\u043d\u0443\u0435\u043c\u043e\u0435" in warnings[0]["suggestion"]


@pytest.mark.unit
class TestCheckDeystvuyushego:
    def test_female_correct(self):
        text = "\u0418\u0432\u0430\u043d\u043e\u0432\u043e\u0439 \u041c\u0430\u0440\u0438\u0438 \u041f\u0435\u0442\u0440\u043e\u0432\u043d\u044b, \u0434\u0435\u0439\u0441\u0442\u0432\u0443\u044e\u0449\u0435\u0439 \u043d\u0430 \u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0438 \u0423\u0441\u0442\u0430\u0432\u0430"
        warnings = check_deystvuyushego(text, 1)
        assert len(warnings) == 0

    def test_female_incorrect(self):
        text = "Ивановой Марии Петровны, действующего на основании Устава"
        warnings = check_deystvuyushego(text, 1)
        assert len(warnings) > 0
        assert "женского" in warnings[0]["suggestion"]

    def test_male_correct(self):
        text = "Иванова Дмитрия Ивановича, действующего на основании Устава"
        warnings = check_deystvuyushego(text, 1)
        assert len(warnings) == 0

    def test_male_incorrect(self):
        text = "Иванова Дмитрия Ивановича, действующей на основании Устава"
        warnings = check_deystvuyushego(text, 1)
        assert len(warnings) > 0
        assert "мужского" in warnings[0]["suggestion"]


@pytest.mark.unit
class TestCheckDocument:
    def test_check_document_with_docx(self, tmp_path):
        from docx import Document
        from scripts.grammar_check import check_document

        doc = Document()
        doc.add_paragraph("ООО «Тест», именуемое в дальнейшем")
        doc.add_paragraph("в лице Генерального директора Иванова")
        doc.add_paragraph("на основании Устава")
        filepath = tmp_path / "test.docx"
        doc.save(str(filepath))

        warnings = check_document(filepath)
        assert isinstance(warnings, list)

    def test_check_document_with_errors(self, tmp_path):
        from docx import Document
        from scripts.grammar_check import check_document

        doc = Document()
        doc.add_paragraph("ООО «Тест», именуемый в дальнейшем")
        doc.add_paragraph("в лице Генеральный директор")
        doc.add_paragraph("на основании Устав")
        filepath = tmp_path / "bad.docx"
        doc.save(str(filepath))

        warnings = check_document(filepath, verbose=True)
        assert len(warnings) >= 1

    def test_check_document_empty(self, tmp_path):
        from docx import Document
        from scripts.grammar_check import check_document

        doc = Document()
        doc.add_paragraph("")
        filepath = tmp_path / "empty.docx"
        doc.save(str(filepath))

        warnings = check_document(filepath)
        assert warnings == []

    def test_check_ip_imenuemiy(self):
        text = "ИП Иванов, именуемый в дальнейшем"
        warnings = check_imenuemoe(text, 1)
        assert len(warnings) == 0

    def test_check_ip_imenuemoe_error(self):
        text = "ИП Иванов, именуемое в дальнейшем"
        warnings = check_imenuemoe(text, 1)
        assert len(warnings) > 0

    def test_no_match_returns_empty(self):
        text = "Обычный текст без юридических конструкций"
        assert check_v_litse(text, 1) == []
        assert check_na_osnovanii(text, 1) == []
        assert check_imenuemoe(text, 1) == []
        assert check_deystvuyushego(text, 1) == []

    def test_is_genitive_function(self):
        from scripts.grammar_check import is_genitive
        assert is_genitive("Устава") is True

    def test_tables_in_docx_checked(self, tmp_path):
        from docx import Document
        from scripts.grammar_check import check_document

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = 'ООО «Пример», именуемый в дальнейшем'
        filepath = tmp_path / "table_check.docx"
        doc.save(str(filepath))

        warnings = check_document(filepath)
        imenuemoe_warnings = [w for w in warnings if w["pattern"] == "именуемое"]
        assert len(imenuemoe_warnings) >= 1

    def test_multiple_errors_in_document(self, tmp_path):
        from docx import Document
        from scripts.grammar_check import check_document

        doc = Document()
        doc.add_paragraph('ООО «Тест», именуемый в дальнейшем «Исполнитель»')
        doc.add_paragraph('ООО «Ещё», именуемый в дальнейшем «Заказчик»')
        filepath = tmp_path / "multi_errors.docx"
        doc.save(str(filepath))

        warnings = check_document(filepath)
        imenuemoe_warnings = [w for w in warnings if w["pattern"] == "именуемое"]
        assert len(imenuemoe_warnings) >= 2


# -----------------------------------------------------------------------
# NEW TESTS: main() function
# -----------------------------------------------------------------------

@pytest.mark.unit
class TestGrammarCheckMain:
    def test_main_with_clean_file(self, tmp_path, capsys):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Просто текст.")
        filepath = tmp_path / "test_main.docx"
        doc.save(str(filepath))

        import sys
        from unittest.mock import patch

        with patch.object(sys, "argv", ["grammar_check", "--file", str(filepath)]):
            from scripts.grammar_check import main
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 0
        captured = capsys.readouterr()
        assert "проблем не найдено" in captured.out

    def test_main_file_not_found(self, tmp_path):
        import sys
        from unittest.mock import patch

        fake = tmp_path / "nonexistent.docx"
        with patch.object(sys, "argv", ["grammar_check", "--file", str(fake)]):
            from scripts.grammar_check import main
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 1

    def test_main_with_warnings_output(self, tmp_path, capsys):
        from docx import Document

        doc = Document()
        doc.add_paragraph('ООО «Тест», именуемый в дальнейшем')
        filepath = tmp_path / "warn_main.docx"
        doc.save(str(filepath))

        import sys
        from unittest.mock import patch

        with patch.object(sys, "argv", ["grammar_check", "--file", str(filepath)]):
            from scripts.grammar_check import main
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 0
            captured = capsys.readouterr()
            assert "предупреждений" in captured.out

    def test_main_verbose_flag(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("на основании Устава")
        filepath = tmp_path / "verbose_main.docx"
        doc.save(str(filepath))

        import sys
        from unittest.mock import patch

        with patch.object(sys, "argv", ["grammar_check", "--file", str(filepath), "--verbose"]):
            from scripts.grammar_check import main
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 0
